"""Resume upload + parse API tests (requires Postgres + MinIO)."""

from __future__ import annotations

import io

import pytest
from pypdf import PdfWriter


def _minimal_pdf_bytes(text: str = "Hello Resume") -> bytes:
    writer = PdfWriter()
    writer.add_blank_page(width=612, height=792)
    # pypdf blank pages have no text; for parse tests we use structure unit tests.
    # Upload/parse integration uses a real text PDF when possible.
    buffer = io.BytesIO()
    writer.write(buffer)
    return buffer.getvalue()


async def _auth_headers(client, unique_email: str) -> dict[str, str]:
    response = await client.post(
        "/api/v1/auth/register",
        json={"email": unique_email, "password": "securepass123", "display_name": "Uploader"},
    )
    assert response.status_code == 201, response.text
    token = response.json()["data"]["access_token"]
    return {"Authorization": f"Bearer {token}"}


@pytest.mark.asyncio
async def test_upload_list_get_resume(client, unique_email):
    headers = await _auth_headers(client, unique_email)
    pdf_bytes = _minimal_pdf_bytes()

    upload = await client.post(
        "/api/v1/resumes",
        headers=headers,
        files={"file": ("resume.pdf", io.BytesIO(pdf_bytes), "application/pdf")},
    )
    assert upload.status_code == 201, upload.text
    resume = upload.json()["data"]
    assert resume["status"] in {"parsing", "extracted", "needs_review", "parse_failed"}
    assert resume["source_mime_type"] == "application/pdf"
    assert resume["source_object_key"]
    assert resume["active_version_id"]

    listed = await client.get("/api/v1/resumes", headers=headers)
    assert listed.status_code == 200
    items = listed.json()["data"]["items"]
    assert len(items) >= 1
    assert items[0]["id"] == resume["id"]

    detail = await client.get(f"/api/v1/resumes/{resume['id']}", headers=headers)
    assert detail.status_code == 200
    assert detail.json()["data"]["id"] == resume["id"]


@pytest.mark.asyncio
async def test_parse_docx_resume(client, unique_email):
    from docx import Document

    headers = await _auth_headers(client, unique_email)
    doc = Document()
    doc.add_paragraph("Alex Example")
    doc.add_paragraph("Backend Engineer")
    doc.add_paragraph("alex@example.com")
    doc.add_paragraph("Experience")
    doc.add_paragraph("Engineer at Acme | Jan 2021 - Present")
    doc.add_paragraph("• Shipped APIs")
    doc.add_paragraph("Skills")
    doc.add_paragraph("Python, FastAPI")
    buffer = io.BytesIO()
    doc.save(buffer)
    docx_bytes = buffer.getvalue()

    upload = await client.post(
        "/api/v1/resumes",
        headers=headers,
        files={
            "file": (
                "resume.docx",
                io.BytesIO(docx_bytes),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert upload.status_code == 201, upload.text
    resume_id = upload.json()["data"]["id"]

    # Background task may have finished; force parse for determinism.
    parsed = await client.post(f"/api/v1/resumes/{resume_id}/parse", headers=headers)
    assert parsed.status_code == 200, parsed.text
    body = parsed.json()["data"]
    assert body["status"] in {"extracted", "needs_review"}
    assert body["content"] is not None
    assert body["content"]["contact"]["email"] == "alex@example.com"
    assert body["content"]["experience"]


@pytest.mark.asyncio
async def test_upload_rejects_unsupported_type(client, unique_email):
    headers = await _auth_headers(client, unique_email)
    response = await client.post(
        "/api/v1/resumes",
        headers=headers,
        files={"file": ("notes.txt", io.BytesIO(b"hello"), "text/plain")},
    )
    assert response.status_code == 415
    assert response.json()["errors"][0]["code"] == "unsupported_type"


@pytest.mark.asyncio
async def test_resumes_require_auth(client):
    response = await client.get("/api/v1/resumes")
    assert response.status_code == 401
